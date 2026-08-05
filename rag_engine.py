import os
import datetime
import traceback
from typing import List, Dict, Any, Tuple

# Direct import to bypass PyTorch/sentence_transformers dependency
from langchain_text_splitters.character import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
from chromadb.utils import embedding_functions

# Document loaders
import docx
import pypdf
import os

APP_NAME = "VCM AI"

APPDATA_DIR = os.path.join(
    os.getenv("LOCALAPPDATA"),
    APP_NAME
)

PERSIST_DIR = os.path.join(
    APPDATA_DIR,
    "vector_db"
)

os.makedirs(PERSIST_DIR, exist_ok=True)


class ONNXEmbeddings:
    """Torch-free, ONNX-based embedding wrapper for Chroma and LangChain."""
    def __init__(self):
        self._ef = embedding_functions.DefaultEmbeddingFunction()

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self._ef(texts)

    def embed_query(self, text: str) -> List[float]:
        return self._ef([text])[0]


class RAGEngine:
    def __init__(self, ollama_model: str = "qwen2.5:7b"):
        self.ollama_model = ollama_model
        self.persist_dir = PERSIST_DIR
        os.makedirs(self.persist_dir, exist_ok=True)
        
        # Initialize Embeddings without PyTorch dependency
        try:
            # First attempt: Ollama embeddings
            self.embeddings = OllamaEmbeddings(model="nomic-embed-text")
            try:
                self.embeddings = OllamaEmbeddings(model="nomic-embed-text")
            except Exception:
                print("[RAGEngine] Using ONNX Runtime Embeddings (torch-free).")
                self.embeddings = ONNXEmbeddings()
            print("[RAGEngine] Using Ollama Embeddings (nomic-embed-text).")
        except Exception:
            # Robust Torch-free ONNX Embeddings fallback
            print("[RAGEngine] Using ONNX Runtime Embeddings (torch-free).")
            self.embeddings = ONNXEmbeddings()

        # Chroma vector stores for Memory and Knowledge
        self.memory_store = Chroma(
            collection_name="vcm_memory",
            embedding_function=self.embeddings,
            persist_directory=os.path.join(self.persist_dir, "memory")
        )
        
        self.knowledge_store = Chroma(
            collection_name="vcm_knowledge",
            embedding_function=self.embeddings,
            persist_directory=os.path.join(self.persist_dir, "knowledge")
        )

    # ===== MEMORY MANAGEMENT =====
    def add_memory(self, memory_text: str) -> str:
        """Stores a fact/memory into the vector database with metadata."""
        if not memory_text.strip():
            return "Nothing to remember."
            
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = Document(
            page_content=memory_text,
            metadata={"timestamp": timestamp, "source": "user_input"}
        )
        self.memory_store.add_documents([doc])
        
        # Also append to memory.txt for backward compatibility
        memory_file = os.path.join(APPDATA_DIR, "memory.txt")
        try:
            with open(memory_file, "a", encoding="utf-8") as f:
                f.write(memory_text + "\n")
        except Exception:
            pass
            
        return "I have saved that into my vector memory."

    def get_relevant_memories(self, query: str, k: int = 3) -> List[str]:
        """Retrieves top-k relevant memories using semantic similarity search."""
        try:
            results = self.memory_store.similarity_search(query, k=k)
            return [doc.page_content for doc in results]
        except Exception:
            traceback.print_exc()
        return []

    def get_all_memories(self) -> List[str]:
        """Returns all raw memories stored in vector store or fallback file."""
        memory_file = os.path.join(APPDATA_DIR, "memory.txt")
        if os.path.exists(memory_file):
            with open(memory_file, "r", encoding="utf-8") as f:
                lines = [line.strip() for line in f.readlines() if line.strip()]
                return lines
        return []

    # ===== DOCUMENT KNOWLEDGE BASE (RAG) =====
    def index_document(self, file_path: str) -> Tuple[bool, str]:
        """Loads, splits, and indexes a file (.txt, .md, .docx, .pdf) into the vector DB."""
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"
            
        ext = os.path.splitext(file_path)[1].lower()
        text = ""

        try:
            if ext in [".txt", ".md", ".py"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif ext == ".docx":
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs if p.text])
            elif ext == ".pdf":
                reader = pypdf.PdfReader(file_path)
                pages = [page.extract_text() for page in reader.pages if page.extract_text()]
                text = "\n".join(pages)
            else:
                return False, f"Unsupported file format: {ext}"

            if not text.strip():
                return False, "Document contains no extractable text."

            # Chunk document
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50
            )
            chunks = splitter.split_text(text)
            
            docs = [
                Document(
                    page_content=chunk,
                    metadata={"source": os.path.basename(file_path), "chunk_id": i}
                )
                for i, chunk in enumerate(chunks)
            ]

            self.knowledge_store.add_documents(docs)
            return True, f"Successfully indexed '{os.path.basename(file_path)}' ({len(chunks)} chunks)."
        except Exception as e:
            return False, f"Error indexing document: {str(e)}"

    def get_relevant_documents(self, query: str, k: int = 4) -> List[Tuple[str, str]]:
        """Retrieves relevant document chunks and their file sources."""
        try:
            results = self.knowledge_store.similarity_search(query, k=k)
            return [(doc.page_content, doc.metadata.get("source", "Unknown")) for doc in results]
        
        except Exception:
            traceback.print_exc()
            return []

    def auto_index_project_docs(self):
        """Automatically indexes project documentation if present."""
        doc_file = os.path.join(os.path.dirname(__file__), "VCMtalker_AI_Project_Documentation.docx")
        if os.path.exists(doc_file):
            self.index_document(doc_file)

    # ===== PROMPT BUILDING =====
    def build_rag_prompt(self, query: str) -> str:

        return f"""You are VCMtalker AI.

            User Question:
            {query}

            Answer naturally.
"""
        
        print("RAG 1")

        memories = self.get_relevant_memories(query, k=3)

        print("RAG 2")

        doc_chunks = self.get_relevant_documents(query, k=3)

        print("RAG 3")

        context_parts = []
        
        if memories:
            context_parts.append("=== RELEVANT USER MEMORY ===")
            for mem in memories:
                context_parts.append(f"- {mem}")

        if doc_chunks:
            context_parts.append("\n=== RELEVANT DOCUMENT KNOWLEDGE BASE ===")
            for chunk, src in doc_chunks:
                context_parts.append(f"Source [{src}]: {chunk}")

        context_str = "\n".join(context_parts) if context_parts else "No specific memory/document context found."

        prompt = f"""You are VCMtalker AI, a helpful, intelligent desktop AI assistant.

{context_str}

User Question: {query}

Instructions:
- Use the relevant memory and document knowledge base provided above if helpful.
- Provide a clear, direct, and conversational answer.
"""
        return prompt
